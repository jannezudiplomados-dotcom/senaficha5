import os
import io
import csv
import re
import uuid
import zipfile
import tempfile
import shutil
from datetime import datetime
from flask import (Blueprint, render_template, request, redirect,
                   url_for, session, flash, send_file, current_app)
from werkzeug.utils import secure_filename
from mysql.connector import Error as MySQLError
import models
from routes.auth import login_required, role_required

documentos_bp = Blueprint('documentos', __name__)


def _permitido(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ('docx', 'xlsx')


@documentos_bp.route('/')
@login_required
def index():
    page = request.args.get('page', 1, type=int)
    per_page = 5
    total = models.contar_plantillas()
    plantillas_paginadas = models.obtener_plantillas_paginadas(page, per_page)
    plantillas_todas = models.listar_plantillas() # For dropdowns
    fichas = models.listar_fichas()
    programas = models.listar_programas()
    total_pages = (total + per_page - 1) // per_page

    return render_template('documentos/index.html',
                           plantillas_paginadas=plantillas_paginadas,
                           plantillas_todas=plantillas_todas,
                           fichas=fichas,
                           programas=programas,
                           page=page, total_pages=total_pages)


# RUTAS DE SUBIDA MOVIDAS A PLANTILLAS


# ---------- HELPERS PARA GENERACION PDF ----------

def _build_context(aprendiz, ficha_info):
    """Build template variable context from aprendiz and ficha data."""
    return {
        'tipo_documento': aprendiz.get('tipo_documento') or '',
        'identificacion': aprendiz.get('identificacion') or '',
        'nombres': aprendiz.get('nombres') or '',
        'apellidos': aprendiz.get('apellidos') or '',
        'correo': aprendiz.get('correo') or '',
        'telefono': aprendiz.get('telefono') or '',
        'direccion': aprendiz.get('direccion') or '',
        'estado': aprendiz.get('estado') or '',
        'numero': ficha_info.get('numero') or '',
        'ficha': ficha_info.get('numero') or '',
        'jornada': ficha_info.get('jornada') or '',
        'fecha_inicio': str(ficha_info.get('fecha_inicio') or ''),
        'fecha_fin': str(ficha_info.get('fecha_fin') or ''),
        'nombre_programa': ficha_info.get('programa_nombre') or '',
        'programa': ficha_info.get('programa_nombre') or '',
        'nombre': ficha_info.get('programa_nombre') or '',
        'codigo_programa': ficha_info.get('programa_codigo') or '',
        'fecha_inicio_programa': str(ficha_info.get('programa_fecha_inicio') or ''),
        'fecha_fin_programa': str(ficha_info.get('programa_fecha_fin') or ''),
        'colegio': ficha_info.get('nombre_colegio') or '',
        'fecha': datetime.now().strftime('%d/%m/%Y'),
    }


def _render_docx(ruta_plantilla, contexto, firma_path, output_path):
    """Render a DOCX template with docxtpl and save."""
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm

    tpl = DocxTemplate(ruta_plantilla)
    ctx = dict(contexto)
    if firma_path and os.path.exists(firma_path):
        ctx['firma'] = InlineImage(tpl, firma_path, width=Mm(40), height=Mm(15))
    else:
        ctx['firma'] = ''
    tpl.render(ctx)
    tpl.save(output_path)


def _add_firma_to_worksheet(ws, firma_path, cell_coordinate):
    """Add a signature image to a worksheet cell with optimal quality.
    The image is embedded at its FULL source resolution but displayed at a
    fixed physical size (~4cm wide) using EMU anchoring.  This way, HiDPI
    800×300 images look extremely crisp at the same visual size as older
    400×150 images."""
    from openpyxl.drawing.image import Image as XlImage
    from openpyxl.utils.units import cm_to_EMU
    from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, AnchorMarker
    from openpyxl.drawing.xdr import XDRPositiveSize2D
    from PIL import Image as PILImage

    try:
        # Read source dimensions for proper aspect ratio
        with PILImage.open(firma_path) as pil_img:
            src_w, src_h = pil_img.size

        # Maximum physical size in cm (fits inside typical Excel cell borders)
        max_w_cm = 3.6
        max_h_cm = 1.35

        # Calculate proportional size to fit within the bounding box
        aspect_ratio = src_w / src_h if src_h > 0 else 1
        
        if aspect_ratio > (max_w_cm / max_h_cm):
            # Image is wider proportionally than the bounding box
            target_w_cm = max_w_cm
            target_h_cm = max_w_cm / aspect_ratio
        else:
            # Image is taller proportionally than the bounding box
            target_h_cm = max_h_cm
            target_w_cm = max_h_cm * aspect_ratio

        img = XlImage(firma_path)

        # Parse cell coordinate to get column and row indices
        from openpyxl.utils import column_index_from_string
        import re as _re
        m = _re.match(r'^([A-Z]+)(\d+)$', cell_coordinate)
        if m:
            col_idx = column_index_from_string(m.group(1)) - 1  # 0-based
            row_idx = int(m.group(2)) - 1  # 0-based
        else:
            col_idx = 0
            row_idx = 0

        # Build OneCellAnchor with EMU dimensions and a small offset (margin)
        # 0.15 cm offset prevents the image background from obscuring the cell's borders
        margin_emu = int(cm_to_EMU(0.15))
        marker = AnchorMarker(col=col_idx, row=row_idx,
                              colOff=margin_emu, rowOff=margin_emu)
        size = XDRPositiveSize2D(
            int(cm_to_EMU(target_w_cm)),
            int(cm_to_EMU(target_h_cm)))
        anchor = OneCellAnchor(_from=marker, ext=size)
        img.anchor = anchor

        ws.add_image(img)
    except Exception:
        pass


def _restore_drawings_from_template(template_path, output_path):
    """Re-inject drawings, media, and shapes from the original template into
    the openpyxl-generated file.  openpyxl drops DrawingML shapes, images
    anchored to drawings, and other non-chart/non-image elements when saving.
    This function MERGES the original drawing anchors (logo, textboxes) with
    any new anchors that openpyxl generated (firma images) so both are preserved.

    Key challenge: openpyxl reuses media names (image1.png, image2.png) for
    new images (firmas), overwriting the template's original media (logo).
    This function renames the generated media to unique names to avoid
    collision, and updates all references accordingly."""
    import zipfile
    from lxml import etree

    # Discover template structure
    tpl_drawing_files = set()
    tpl_media_files = set()
    tpl_drawing_rels = set()

    with zipfile.ZipFile(template_path, 'r') as zt:
        for name in zt.namelist():
            lower = name.lower()
            if lower.startswith('xl/drawings/') and not lower.endswith('.rels'):
                tpl_drawing_files.add(name)
            elif lower.startswith('xl/drawings/_rels/'):
                tpl_drawing_rels.add(name)
            elif lower.startswith('xl/media/'):
                tpl_media_files.add(name)

    if not tpl_drawing_files and not tpl_media_files:
        return  # Template has no drawings; nothing to restore

    tmp_output = output_path + '.tmp.xlsx'

    with zipfile.ZipFile(template_path, 'r') as zt, \
         zipfile.ZipFile(output_path, 'r') as zo, \
         zipfile.ZipFile(tmp_output, 'w', zipfile.ZIP_DEFLATED) as zout:

        generated_names = set(zo.namelist())
        template_names = set(zt.namelist())

        # Identify generated media files (firma images from openpyxl)
        gen_media_files = [n for n in generated_names
                           if n.lower().startswith('xl/media/')]

        # Build rename map: generated media -> unique names that don't
        # collide with template media
        # Find max image number in template media
        import re as _re
        max_img_num = 0
        for name in tpl_media_files:
            m = _re.search(r'image(\d+)', name)
            if m:
                max_img_num = max(max_img_num, int(m.group(1)))

        # media_rename_map: old_zip_path -> new_zip_path
        # media_target_rename: old_target_in_rels -> new_target_in_rels
        media_rename_map = {}
        media_target_rename = {}
        for gen_name in gen_media_files:
            basename = os.path.basename(gen_name)
            ext = os.path.splitext(basename)[1]
            max_img_num += 1
            new_basename = f'firma_image{max_img_num}{ext}'
            new_path = f'xl/media/{new_basename}'
            media_rename_map[gen_name] = new_path
            # openpyxl uses absolute targets like /xl/media/image1.png
            media_target_rename[f'/xl/media/{basename}'] = f'../media/{new_basename}'
            # Also handle relative targets
            media_target_rename[f'../media/{basename}'] = f'../media/{new_basename}'

        # ----- 1. Copy non-drawing/media/content-types from generated -----
        for name in zo.namelist():
            lower = name.lower()
            if lower.startswith('xl/drawings/') or lower.startswith('xl/media/'):
                continue
            if lower == '[content_types].xml':
                continue
            zout.writestr(name, zo.read(name))

        # ----- 2. Copy ALL template media (logo, decorative images) -----
        for name in tpl_media_files:
            zout.writestr(name, zt.read(name))

        # Copy renamed generated media (firma images)
        for old_name, new_name in media_rename_map.items():
            zout.writestr(new_name, zo.read(old_name))

        # ----- 3. Merge drawings per sheet -----
        # Process each drawing file (drawing1.xml, drawing2.xml, etc.)
        # Template's sheet1 has the main drawings (logo, textboxes)
        # openpyxl may generate drawing1.xml for sheet1, drawing2.xml for sheet2, etc.

        all_drawing_names = set()
        for name in tpl_drawing_files:
            all_drawing_names.add(name)
        for name in generated_names:
            if name.lower().startswith('xl/drawings/') and not name.lower().endswith('.rels'):
                all_drawing_names.add(name)

        for drawing_name in sorted(all_drawing_names):
            rels_name = drawing_name.replace('xl/drawings/',
                                             'xl/drawings/_rels/') + '.rels'

            tpl_has = drawing_name in template_names
            gen_has = drawing_name in generated_names

            if tpl_has and gen_has:
                # MERGE: template anchors + generated anchors
                tpl_drawing = etree.fromstring(zt.read(drawing_name))
                gen_drawing = etree.fromstring(zo.read(drawing_name))

                # Process rels: find max rId in template, remap generated rIds
                tpl_rels_doc = None
                max_tpl_rid = 0
                if rels_name in template_names:
                    tpl_rels_doc = etree.fromstring(zt.read(rels_name))
                    for rel in tpl_rels_doc:
                        rid_str = rel.get('Id', 'rId0')
                        try:
                            rid_num = int(rid_str.replace('rId', ''))
                            max_tpl_rid = max(max_tpl_rid, rid_num)
                        except ValueError:
                            pass

                gen_rels_doc = None
                rid_map = {}
                if rels_name in generated_names:
                    gen_rels_doc = etree.fromstring(zo.read(rels_name))
                    for rel in gen_rels_doc:
                        old_rid = rel.get('Id')
                        try:
                            old_num = int(old_rid.replace('rId', ''))
                        except ValueError:
                            old_num = max_tpl_rid + 1
                        new_num = max_tpl_rid + old_num
                        new_rid = f'rId{new_num}'
                        rid_map[old_rid] = new_rid
                        rel.set('Id', new_rid)
                        # Also rename the Target to use the new media name
                        old_target = rel.get('Target', '')
                        if old_target in media_target_rename:
                            rel.set('Target', media_target_rename[old_target])

                # Remap rId references inside generated drawing anchors
                r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

                def _remap_rids(elem):
                    for child in elem.iter():
                        for attr in [f'{{{r_ns}}}embed', f'{{{r_ns}}}link']:
                            val = child.get(attr)
                            if val and val in rid_map:
                                child.set(attr, rid_map[val])

                # Find max element id in template drawing
                max_id = 0
                for elem in tpl_drawing.iter():
                    id_val = elem.get('id')
                    if id_val and id_val.isdigit():
                        max_id = max(max_id, int(id_val))

                # Add generated anchors to template drawing
                for anchor in list(gen_drawing):
                    _remap_rids(anchor)
                    for elem in anchor.iter():
                        id_val = elem.get('id')
                        if id_val and id_val.isdigit():
                            max_id += 1
                            elem.set('id', str(max_id))
                    tpl_drawing.append(anchor)

                # Write merged drawing
                zout.writestr(drawing_name, etree.tostring(
                    tpl_drawing, xml_declaration=True,
                    encoding='UTF-8', standalone=True))

                # Merge rels
                merged_rels = tpl_rels_doc if tpl_rels_doc is not None else gen_rels_doc
                if tpl_rels_doc is not None and gen_rels_doc is not None:
                    for rel in gen_rels_doc:
                        tpl_rels_doc.append(rel)
                    merged_rels = tpl_rels_doc

                if merged_rels is not None:
                    zout.writestr(rels_name, etree.tostring(
                        merged_rels, xml_declaration=True,
                        encoding='UTF-8', standalone=True))

            elif tpl_has:
                # Only template has this drawing
                zout.writestr(drawing_name, zt.read(drawing_name))
                if rels_name in template_names:
                    zout.writestr(rels_name, zt.read(rels_name))

            elif gen_has:
                # Only generated has this drawing (firma images on a sheet
                # that didn't have drawings in the template)
                # Need to rename media targets in the rels
                gen_drawing_xml = zo.read(drawing_name)
                zout.writestr(drawing_name, gen_drawing_xml)

                if rels_name in generated_names:
                    gen_rels_doc = etree.fromstring(zo.read(rels_name))
                    for rel in gen_rels_doc:
                        old_target = rel.get('Target', '')
                        if old_target in media_target_rename:
                            rel.set('Target', media_target_rename[old_target])
                    zout.writestr(rels_name, etree.tostring(
                        gen_rels_doc, xml_declaration=True,
                        encoding='UTF-8', standalone=True))

        # ----- 4. Rebuild [Content_Types].xml -----
        ct_gen = etree.fromstring(zo.read('[Content_Types].xml'))
        ct_tpl = etree.fromstring(zt.read('[Content_Types].xml'))

        existing_parts = set()
        existing_extensions = set()
        for child in ct_gen:
            pn = child.get('PartName')
            ext = child.get('Extension')
            if pn:
                existing_parts.add(pn)
            if ext:
                existing_extensions.add(ext)

        # Add missing Override entries from template
        for child in ct_tpl:
            pn = child.get('PartName')
            if pn and pn not in existing_parts:
                ct_gen.append(child)
                existing_parts.add(pn)

        # Ensure png Default type exists
        if 'png' not in existing_extensions:
            ns_prefix = ct_gen.tag.split('}')[0] + '}' if '}' in ct_gen.tag else ''
            tag_name = f'{ns_prefix}Default' if ns_prefix else 'Default'
            ext_elem = etree.SubElement(ct_gen, tag_name)
            ext_elem.set('Extension', 'png')
            ext_elem.set('ContentType', 'image/png')

        zout.writestr('[Content_Types].xml', etree.tostring(
            ct_gen, xml_declaration=True, encoding='UTF-8', standalone=True))

    # Replace original output with patched version
    os.replace(tmp_output, output_path)


def _render_xlsx(ruta_plantilla, contexto, firma_path, output_path):
    """Render an Excel template replacing {{ variable }} in cells."""
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as XlImage

    wb = load_workbook(ruta_plantilla)
    # Acepta {{ variable }}, {{variable}}, {{ variable } (con llave faltante), etc.
    # Evita comerse los espacios que están FUERA de las llaves.
    pattern = re.compile(r'\{\{\s*(\w+)\s*\}*')

    for ws in wb.worksheets:
        firma_cells = []
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if pattern.search(cell.value):
                        has_firma = bool(re.search(
                            r'\{\{\s*firma[}\s]*', cell.value))
                        if has_firma:
                            firma_cells.append(cell)

                        def _repl(match):
                            var = match.group(1)
                            if var == 'firma':
                                return ''
                            return str(contexto.get(var, match.group(0)))

                        cell.value = pattern.sub(_repl, cell.value)

        if firma_path and os.path.exists(firma_path) and firma_cells:
            for cell in firma_cells:
                _add_firma_to_worksheet(ws, firma_path, cell.coordinate)

    wb.save(output_path)
    # Restaurar drawings/imágenes originales del template (logo, shapes, etc.)
    _restore_drawings_from_template(ruta_plantilla, output_path)


def _render_xlsx_general(ruta_plantilla, aprendices, f_info, output_path):
    """Render an Excel template for ALL aprendices in a single file (Listado)."""
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as XlImage
    import copy

    wb = load_workbook(ruta_plantilla)
    # Acepta {{ variable }}, {{variable}}, {{ variable } (con llave faltante), etc.
    # Evita comerse los espacios que están FUERA de las llaves.
    pattern = re.compile(r'\{\{\s*(\w+)\s*\}*')

    for ws in wb.worksheets:
        template_row_idx = None
        template_cells = []

        # Encontrar la primera fila que contiene variables
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and pattern.search(cell.value):
                    template_row_idx = cell.row
                    break
            if template_row_idx:
                break

        if not template_row_idx:
            continue

        # Guardar estilos y valores de la fila plantilla
        for col_idx in range(1, ws.max_column + 1):
            cell = ws.cell(row=template_row_idx, column=col_idx)
            template_cells.append({
                'value': cell.value,
                'font': copy.copy(cell.font),
                'border': copy.copy(cell.border),
                'fill': copy.copy(cell.fill),
                'number_format': cell.number_format,
                'alignment': copy.copy(cell.alignment)
            })

        current_row = template_row_idx
        for i, a in enumerate(aprendices):
            if i > 0:
                ws.insert_rows(current_row)

            ctx = _build_context(a, f_info)
            firma_path = None
            if a.get('firma'):
                from flask import current_app
                firma_path = os.path.join(current_app.config['FIRMAS_FOLDER'], a['firma'])
            
            has_firma_in_row = False
            firma_col = None

            for col_idx, t_cell in enumerate(template_cells, start=1):
                cell = ws.cell(row=current_row, column=col_idx)

                # Aplicar estilos copiados
                cell.font = copy.copy(t_cell['font'])
                cell.border = copy.copy(t_cell['border'])
                cell.fill = copy.copy(t_cell['fill'])
                cell.number_format = t_cell['number_format']
                cell.alignment = copy.copy(t_cell['alignment'])
                
                val = t_cell['value']
                if val and isinstance(val, str) and pattern.search(val):
                    if re.search(r'\{\{\s*firma[}\s]*', val):
                        has_firma_in_row = True
                        firma_col = col_idx

                    def _repl(match):
                        var = match.group(1)
                        if var == 'firma':
                            return ''
                        return str(ctx.get(var, match.group(0)))

                    cell.value = pattern.sub(_repl, val)
                else:
                    cell.value = val

            if has_firma_in_row and firma_path and os.path.exists(firma_path):
                target_cell = ws.cell(row=current_row, column=firma_col)
                _add_firma_to_worksheet(ws, firma_path, target_cell.coordinate)

            current_row += 1

    wb.save(output_path)
    # Restaurar drawings/imágenes originales del template (logo, shapes, etc.)
    _restore_drawings_from_template(ruta_plantilla, output_path)


def _convert_to_pdf(input_path, output_path):
    """Convert DOCX or XLSX to PDF via subprocess (avoids COM hang in Flask)."""
    import subprocess
    import sys

    script = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                          'convert_to_pdf.py')
    abs_in = os.path.abspath(input_path)
    abs_out = os.path.abspath(output_path)

    result = subprocess.run(
        [sys.executable, script, abs_in, abs_out],
        capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        error_msg = result.stderr.strip() or 'Error desconocido en conversion'
        raise RuntimeError(f'Error al convertir a PDF: {error_msg}')

    if not os.path.exists(abs_out):
        raise RuntimeError('El archivo PDF no fue generado')


@documentos_bp.route('/generar-por-ficha', methods=['POST'])
@login_required
def generar_por_ficha():
    plantilla_id = request.form.get('plantilla_id')
    ficha_id = request.form.get('ficha_id')
    plantilla = models.obtener_plantilla(int(plantilla_id)) if plantilla_id else None
    ficha = models.obtener_ficha(int(ficha_id)) if ficha_id else None

    if not plantilla or not ficha:
        flash('Selecciona una plantilla y una ficha.', 'danger')
        return redirect(url_for('documentos.index'))

    aprendices = models.usuarios_por_ficha(int(ficha_id))
    if not aprendices:
        flash('La ficha seleccionada no tiene aprendices.', 'warning')
        return redirect(url_for('documentos.index'))

    ruta_plantilla = os.path.join(
        current_app.config['UPLOAD_FOLDER'], plantilla['archivo'])
    if not os.path.exists(ruta_plantilla):
        flash('El archivo de la plantilla no existe.', 'danger')
        return redirect(url_for('documentos.index'))

    ext = os.path.splitext(plantilla['archivo'])[1].lower()
    tmp_dir = tempfile.mkdtemp(
        dir=current_app.config['GENERADOS_FOLDER'])

    try:
        pdf_files = []
        errores = []

        for aprendiz in aprendices:
            contexto = _build_context(aprendiz, ficha)
            firma_path = None
            if aprendiz.get('firma'):
                fp = os.path.join(
                    current_app.config['FIRMAS_FOLDER'], aprendiz['firma'])
                if os.path.exists(fp):
                    firma_path = fp

            safe = f"{aprendiz['identificacion']}_{uuid.uuid4().hex[:6]}"
            rendered = os.path.join(tmp_dir, f"{safe}{ext}")

            try:
                if ext == '.docx':
                    _render_docx(ruta_plantilla, contexto,
                                 firma_path, rendered)
                elif ext == '.xlsx':
                    _render_xlsx(ruta_plantilla, contexto,
                                 firma_path, rendered)
                else:
                    continue

                pdf_name = (f"{aprendiz['apellidos']}_"
                            f"{aprendiz['nombres']}_"
                            f"{aprendiz['identificacion']}.pdf")
                pdf_name = "".join(
                    c for c in pdf_name
                    if c.isalnum() or c in ('_', '-', '.', ' '))
                pdf_path = os.path.join(tmp_dir, pdf_name)

                _convert_to_pdf(rendered, pdf_path)

                if os.path.exists(pdf_path):
                    pdf_files.append((pdf_name, pdf_path))
            except Exception as e:
                errores.append(
                    f"{aprendiz['nombres']} {aprendiz['apellidos']}: {e}")
                current_app.logger.warning(
                    'Error PDF %s: %s', aprendiz['identificacion'], e)
            finally:
                if os.path.exists(rendered):
                    try:
                        os.remove(rendered)
                    except OSError:
                        pass

        if not pdf_files:
            msg = 'No se genero ningun PDF.'
            if errores:
                msg += ' ' + '; '.join(errores[:3])
            flash(msg, 'danger')
            return redirect(url_for('documentos.index'))

        if errores:
            flash(f'{len(pdf_files)} PDFs generados, '
                  f'{len(errores)} error(es).', 'warning')

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for name, path in pdf_files:
                zf.write(path, name)
        zip_buf.seek(0)

        models.registrar_log(
            session.get('admin_id'), session.get('admin_username'),
            'GENERAR', 'documentos_ficha', ficha['id'],
            f"Ficha {ficha['numero']} - {len(pdf_files)} PDFs",
            request.remote_addr)

        return send_file(
            zip_buf, as_attachment=True,
            download_name=f"documentos_ficha_{ficha['numero']}.zip",
            mimetype='application/zip')

    except Exception as e:
        current_app.logger.exception('Error generando documentos por ficha')
        flash('Error al generar los documentos. Consulta con el administrador.', 'danger')
        return redirect(url_for('documentos.index'))

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


@documentos_bp.route('/generar-general', methods=['POST'])
@login_required
def generar_general():
    plantilla_id = request.form.get('plantilla_id')
    ficha_id = request.form.get('ficha_id')

    if not plantilla_id or not ficha_id:
        flash('Seleccione plantilla y ficha', 'danger')
        return redirect(url_for('documentos.index'))

    plantilla = models.obtener_plantilla(plantilla_id)
    ficha_info = models.obtener_ficha(ficha_id)
    if not plantilla or not ficha_info:
        flash('Datos inválidos.', 'danger')
        return redirect(url_for('documentos.index'))

    ext = plantilla['archivo'].split('.')[-1].lower()
    if ext != 'xlsx':
        flash('La plantilla general debe ser un archivo de Excel (.xlsx).', 'danger')
        return redirect(url_for('documentos.index'))

    ruta_plantilla = os.path.join(current_app.config['UPLOAD_FOLDER'], plantilla['archivo'])
    aprendices = models.usuarios_por_ficha(ficha_id)

    if not aprendices:
        flash('La ficha no tiene aprendices registrados.', 'warning')
        return redirect(url_for('documentos.index'))

    temp_dir = tempfile.mkdtemp(dir=current_app.config['GENERADOS_FOLDER'])
    try:
        temp_xlsx = os.path.join(temp_dir, f"listado_ficha_{ficha_info['numero']}.xlsx")
        temp_pdf = os.path.join(temp_dir, f"Listado_General_Ficha_{ficha_info['numero']}.pdf")

        # Generar Excel con todos los aprendices
        _render_xlsx_general(ruta_plantilla, aprendices, ficha_info, temp_xlsx)

        # Convertir a PDF
        _convert_to_pdf(temp_xlsx, temp_pdf)

        if not os.path.exists(temp_pdf):
            raise Exception('El archivo PDF final no se generó correctamente.')

        models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                             'GENERAR_GENERAL', 'documentos', None,
                             f"Ficha {ficha_info['numero']}", request.remote_addr)

        with open(temp_pdf, 'rb') as f:
            pdf_data = f.read()

        buffer = io.BytesIO(pdf_data)
        buffer.seek(0)
        
        return send_file(buffer, as_attachment=True,
                         download_name=f"Listado_General_Ficha_{ficha_info['numero']}.pdf",
                         mimetype='application/pdf')

    except Exception as e:
        current_app.logger.error(f"Error generando documento general: {e}")
        flash('Error al generar el documento general. Consulta con el administrador.', 'danger')
        return redirect(url_for('documentos.index'))
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


# ---------- LISTADO COMPLETO DE UNA FICHA ----------
@documentos_bp.route('/ficha-listado', methods=['POST'])
@login_required
def ficha_listado():
    ficha_id = request.form.get('ficha_id')
    formato = request.form.get('formato', 'word')
    if not ficha_id:
        flash('Selecciona una ficha.', 'danger')
        return redirect(url_for('documentos.index'))
    f = models.obtener_ficha(int(ficha_id))
    if not f:
        flash('Ficha no encontrada.', 'danger')
        return redirect(url_for('documentos.index'))
    aprendices = models.usuarios_por_ficha(int(ficha_id))
    models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                         'GENERAR', 'ficha', f['id'],
                         f"Listado {f['numero']} ({formato})", request.remote_addr)
    if formato == 'excel':
        return _listado_excel(f, aprendices)
    return _listado_word(f, aprendices)


def _listado_word(f, aprendices):
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    titulo = doc.add_heading(f"Listado de aprendices - Ficha {f['numero']}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Programa: {f['programa_nombre']}")
    doc.add_paragraph(
        f"Jornada: {f['jornada']}    Inicio: {f.get('fecha_inicio') or '-'}    "
        f"Finalizacion: {f.get('fecha_fin') or '-'}")
    doc.add_paragraph(f"Total aprendices: {len(aprendices)}")

    encabezados = ['Tipo Doc', 'No. Identificacion', 'Nombres', 'Apellidos', 'Direccion',
                   'Correo', 'Telefono', 'Inicio', 'Finalizacion', 'Firma']
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Table Grid'
    hdr = tabla.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = h
        for parrafo in hdr[i].paragraphs:
            for run in parrafo.runs:
                run.font.bold = True
                run.font.size = Pt(8)

    for a in aprendices:
        celdas = tabla.add_row().cells
        valores = [
            a.get('tipo_documento') or '',
            str(a.get('identificacion') or ''),
            a.get('nombres') or '',
            a.get('apellidos') or '',
            a.get('direccion') or '',
            a.get('correo') or '',
            a.get('telefono') or '',
            str(f.get('fecha_inicio') or ''),
            str(f.get('fecha_fin') or ''),
        ]
        for i, v in enumerate(valores):
            celdas[i].text = v
            for parrafo in celdas[i].paragraphs:
                for run in parrafo.runs:
                    run.font.size = Pt(8)
        if a.get('firma'):
            ruta = os.path.join(current_app.config['FIRMAS_FOLDER'], a['firma'])
            if os.path.exists(ruta):
                run = celdas[9].paragraphs[0].add_run()
                try:
                    run.add_picture(ruta, width=Mm(22))
                except Exception:
                    celdas[9].text = 'Firma'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name=f"ficha_{f['numero']}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def _listado_excel(f, aprendices):
    from openpyxl import Workbook
    from openpyxl.styles import Font
    wb = Workbook()
    ws = wb.active
    ws.title = f"Ficha {f['numero']}"[:31]
    ws.append([f"Ficha {f['numero']} - {f['programa_nombre']}"])
    ws.append([f"Jornada: {f['jornada']}  Inicio: {f.get('fecha_inicio') or '-'}  "
               f"Finalizacion: {f.get('fecha_fin') or '-'}"])
    ws.append([])
    encabezados = ['Tipo Doc', 'No. Identificacion', 'Nombres', 'Apellidos', 'Direccion',
                   'Correo', 'Telefono', 'Inicio', 'Finalizacion', 'Firma']
    ws.append(encabezados)
    for c in ws[4]:
        c.font = Font(bold=True)
    for a in aprendices:
        ws.append([
            a.get('tipo_documento') or '',
            a.get('identificacion') or '',
            a.get('nombres') or '',
            a.get('apellidos') or '',
            a.get('direccion') or '',
            a.get('correo') or '',
            a.get('telefono') or '',
            str(f.get('fecha_inicio') or ''),
            str(f.get('fecha_fin') or ''),
            'Si' if a.get('firma') else 'No',
        ])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name=f"ficha_{f['numero']}.xlsx",
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@documentos_bp.route('/exportar/excel')
@login_required
def exportar_excel():
    from openpyxl import Workbook
    aprendices = models.obtener_usuarios_paginados(1, 100000)
    wb = Workbook()
    ws = wb.active
    ws.title = 'Aprendices'
    ws.append(['ID', 'Identificacion', 'Tipo', 'Nombres', 'Apellidos', 'Direccion',
               'Correo', 'Telefono', 'Ficha', 'Programa', 'Estado'])
    for a in aprendices:
        ws.append([a['id'], a['identificacion'], a['tipo_documento'], a['nombres'],
                   a['apellidos'], a.get('direccion') or '', a.get('correo') or '',
                   a.get('telefono') or '', a.get('ficha_numero') or '',
                   a.get('programa_nombre') or '', a['estado']])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                         'EXPORTAR', 'usuario', None, 'Excel', request.remote_addr)
    return send_file(buffer, as_attachment=True, download_name='aprendices.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@documentos_bp.route('/exportar/csv')
@login_required
def exportar_csv():
    aprendices = models.obtener_usuarios_paginados(1, 100000)
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['ID', 'Identificacion', 'Tipo', 'Nombres', 'Apellidos', 'Direccion',
                     'Correo', 'Telefono', 'Ficha', 'Programa', 'Estado'])
    for a in aprendices:
        writer.writerow([a['id'], a['identificacion'], a['tipo_documento'], a['nombres'],
                         a['apellidos'], a.get('direccion') or '', a.get('correo') or '',
                         a.get('telefono') or '', a.get('ficha_numero') or '',
                         a.get('programa_nombre') or '', a['estado']])
    mem = io.BytesIO(buffer.getvalue().encode('utf-8-sig'))
    mem.seek(0)
    models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                         'EXPORTAR', 'usuario', None, 'CSV', request.remote_addr)
    return send_file(mem, as_attachment=True, download_name='aprendices.csv', mimetype='text/csv')


# ---------- GESTION INDIVIDUAL ----------

@documentos_bp.route('/gestion-individual')
@login_required
def gestion_individual():
    plantillas_todas = models.listar_plantillas()
    fichas = models.listar_fichas()
    programas = models.listar_programas()
    return render_template('documentos/gestion_individual.html',
                           plantillas_todas=plantillas_todas,
                           fichas=fichas,
                           programas=programas)

@documentos_bp.route('/api/aprendices/<int:ficha_id>')
@login_required
def api_aprendices(ficha_id):
    from flask import jsonify
    aprendices = models.usuarios_por_ficha(ficha_id)
    return jsonify([{
        'id': a['id'],
        'nombres': a['nombres'],
        'apellidos': a['apellidos'],
        'identificacion': a['identificacion']
    } for a in aprendices])

@documentos_bp.route('/generar-individual', methods=['POST'])
@login_required
def generar_individual():
    plantilla_id = request.form.get('plantilla_id')
    ficha_id = request.form.get('ficha_id')
    aprendiz_id = request.form.get('aprendiz_id')

    if not plantilla_id or not ficha_id or not aprendiz_id:
        flash('Datos incompletos. Selecciona plantilla, ficha y aprendiz.', 'danger')
        return redirect(url_for('documentos.gestion_individual'))

    plantilla = models.obtener_plantilla(int(plantilla_id))
    ficha = models.obtener_ficha(int(ficha_id))
    aprendiz = models.obtener_usuario(int(aprendiz_id))

    if not plantilla or not ficha or not aprendiz:
        flash('Datos no encontrados en el sistema.', 'danger')
        return redirect(url_for('documentos.gestion_individual'))

    ruta_plantilla = os.path.join(current_app.config['UPLOAD_FOLDER'], plantilla['archivo'])
    if not os.path.exists(ruta_plantilla):
        flash('El archivo de la plantilla no existe.', 'danger')
        return redirect(url_for('documentos.gestion_individual'))

    ext = os.path.splitext(plantilla['archivo'])[1].lower()
    tmp_dir = tempfile.mkdtemp(dir=current_app.config['GENERADOS_FOLDER'])

    try:
        contexto = _build_context(aprendiz, ficha)
        firma_path = None
        if aprendiz.get('firma'):
            fp = os.path.join(current_app.config['FIRMAS_FOLDER'], aprendiz['firma'])
            if os.path.exists(fp):
                firma_path = fp

        safe = f"{aprendiz['identificacion']}_{uuid.uuid4().hex[:6]}"
        rendered = os.path.join(tmp_dir, f"{safe}{ext}")

        if ext == '.docx':
            _render_docx(ruta_plantilla, contexto, firma_path, rendered)
        elif ext == '.xlsx':
            _render_xlsx(ruta_plantilla, contexto, firma_path, rendered)
        else:
            flash('Formato de plantilla no soportado.', 'danger')
            return redirect(url_for('documentos.gestion_individual'))

        pdf_name = (f"{aprendiz['apellidos']}_"
                    f"{aprendiz['nombres']}_"
                    f"{aprendiz['identificacion']}.pdf")
        pdf_name = "".join(c for c in pdf_name if c.isalnum() or c in ('_', '-', '.', ' '))
        pdf_path = os.path.join(tmp_dir, pdf_name)

        _convert_to_pdf(rendered, pdf_path)

        if not os.path.exists(pdf_path):
            raise RuntimeError('El PDF individual no fue generado correctamente.')

        models.registrar_log(
            session.get('admin_id'), session.get('admin_username'),
            'GENERAR', 'documentos_individual', aprendiz['id'],
            f"Ficha {ficha['numero']} - PDF de {aprendiz['nombres']} {aprendiz['apellidos']}",
            request.remote_addr)

        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()

        buffer = io.BytesIO(pdf_data)
        buffer.seek(0)
        
        return send_file(buffer, as_attachment=True,
                         download_name=pdf_name,
                         mimetype='application/pdf')

    except Exception as e:
        current_app.logger.exception('Error generando documento individual')
        flash('Error al generar el documento individual. Consulta con el administrador.', 'danger')
        return redirect(url_for('documentos.gestion_individual'))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------- GESTION GRUPAL ----------

@documentos_bp.route('/gestion-grupal')
@login_required
def gestion_grupal():
    plantillas_todas = models.listar_plantillas()
    fichas = models.listar_fichas()
    programas = models.listar_programas()
    return render_template('documentos/gestion_grupal.html',
                           plantillas_todas=plantillas_todas,
                           fichas=fichas,
                           programas=programas)

@documentos_bp.route('/generar-grupal', methods=['POST'])
@login_required
def generar_grupal():
    plantilla_id = request.form.get('plantilla_id')
    ficha_id = request.form.get('ficha_id')
    aprendices_ids = request.form.getlist('aprendices_ids')

    if not plantilla_id or not ficha_id or not aprendices_ids:
        flash('Datos incompletos. Selecciona plantilla, ficha y al menos un aprendiz.', 'danger')
        return redirect(url_for('documentos.gestion_grupal'))

    plantilla = models.obtener_plantilla(int(plantilla_id))
    ficha_info = models.obtener_ficha(int(ficha_id))
    if not plantilla or not ficha_info:
        flash('Datos inválidos.', 'danger')
        return redirect(url_for('documentos.gestion_grupal'))

    ext = plantilla['archivo'].split('.')[-1].lower()

    ruta_plantilla = os.path.join(current_app.config['UPLOAD_FOLDER'], plantilla['archivo'])
    
    # Obtener todos los aprendices y filtrar los seleccionados
    todos_aprendices = models.usuarios_por_ficha(int(ficha_id))
    aprendices = [a for a in todos_aprendices if str(a['id']) in aprendices_ids]

    if not aprendices:
        flash('No se encontraron los aprendices seleccionados.', 'warning')
        return redirect(url_for('documentos.gestion_grupal'))

    temp_dir = tempfile.mkdtemp(dir=current_app.config['GENERADOS_FOLDER'])
    try:
        if ext == 'xlsx':
            temp_xlsx = os.path.join(temp_dir, f"listado_grupal_ficha_{ficha_info['numero']}.xlsx")
            temp_pdf = os.path.join(temp_dir, f"Listado_Grupal_Ficha_{ficha_info['numero']}.pdf")

            # Generar Excel con aprendices seleccionados
            _render_xlsx_general(ruta_plantilla, aprendices, ficha_info, temp_xlsx)

            # Convertir a PDF
            _convert_to_pdf(temp_xlsx, temp_pdf)

            if not os.path.exists(temp_pdf):
                raise Exception('El archivo PDF final no se generó correctamente.')

            models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                                 'GENERAR_GRUPAL', 'documentos', None,
                                 f"Ficha {ficha_info['numero']} - {len(aprendices)} aprendices", request.remote_addr)

            with open(temp_pdf, 'rb') as f:
                pdf_data = f.read()

            buffer = io.BytesIO(pdf_data)
            buffer.seek(0)
            
            return send_file(buffer, as_attachment=True,
                             download_name=f"Listado_Grupal_Ficha_{ficha_info['numero']}.pdf",
                             mimetype='application/pdf')

        elif ext == 'docx':
            from docxtpl import DocxTemplate, InlineImage
            from docx.shared import Mm
            
            tpl = DocxTemplate(ruta_plantilla)
            
            # Construir la lista de aprendices con sus imágenes preparadas
            lista_aprendices = []
            for a in aprendices:
                ctx = _build_context(a, ficha_info)
                firma_path = os.path.join(current_app.config['FIRMAS_FOLDER'], a['firma']) if a.get('firma') else None
                if firma_path and os.path.exists(firma_path):
                    ctx['firma'] = InlineImage(tpl, firma_path, width=Mm(40), height=Mm(15))
                else:
                    ctx['firma'] = ''
                lista_aprendices.append(ctx)
                
            # Construir contexto general con la lista completa
            contexto_general = lista_aprendices[0].copy() if lista_aprendices else {}
            contexto_general['aprendices'] = lista_aprendices
            
            # Generar el documento final único
            temp_docx = os.path.join(temp_dir, f"Listado_Grupal_Ficha_{ficha_info['numero']}.docx")
            temp_pdf = os.path.join(temp_dir, f"Listado_Grupal_Ficha_{ficha_info['numero']}.pdf")
            
            tpl.render(contexto_general)
            tpl.save(temp_docx)
            
            _convert_to_pdf(temp_docx, temp_pdf)
            
            if not os.path.exists(temp_pdf):
                raise Exception('El archivo PDF final no se generó correctamente.')
                
            models.registrar_log(session.get('admin_id'), session.get('admin_username'),
                                 'GENERAR_GRUPAL', 'documentos', None,
                                 f"Ficha {ficha_info['numero']} - {len(aprendices)} aprendices (Word único)", request.remote_addr)

            with open(temp_pdf, 'rb') as f:
                pdf_data = f.read()

            buffer = io.BytesIO(pdf_data)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True,
                             download_name=f"Listado_Grupal_Ficha_{ficha_info['numero']}.pdf",
                             mimetype='application/pdf')
                             
        else:
            flash('Formato de plantilla no soportado (.xlsx o .docx).', 'danger')
            return redirect(url_for('documentos.gestion_grupal'))

    except Exception as e:
        current_app.logger.error(f"Error generando documento grupal: {e}")
        flash('Error al generar el documento grupal. Consulta con el administrador.', 'danger')
        return redirect(url_for('documentos.gestion_grupal'))
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
